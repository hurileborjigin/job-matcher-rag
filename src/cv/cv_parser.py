"""
CV Parser - Extract text from various CV file formats
Supports: PDF, DOCX, TXT, and images (with OCR)
"""
import io
import logging
from typing import Optional, Dict
from pathlib import Path

# PDF parsing - using pypdf (not PyPDF2)
try:
    from pypdf import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    try:
        # Fallback to PyPDF2 if available
        from PyPDF2 import PdfReader
        PDF_AVAILABLE = True
    except ImportError:
        PDF_AVAILABLE = False
        logging.warning("pypdf not available. PDF parsing disabled.")

# DOCX parsing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available. DOCX parsing disabled.")

# Image OCR (optional)
try:
    from PIL import Image
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    logging.warning("OCR not available. Image parsing disabled.")

logger = logging.getLogger(__name__)


class CVParser:
    """Parse CV files and extract text content."""
    
    SUPPORTED_FORMATS = {
        'txt': 'Plain Text',
        'pdf': 'PDF Document',
        'docx': 'Word Document',
        'doc': 'Word Document (Legacy)',
        'png': 'Image (PNG)',
        'jpg': 'Image (JPEG)',
        'jpeg': 'Image (JPEG)'
    }
    
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB
    
    def __init__(self):
        """Initialize CV parser."""
        self.supported_formats = self._get_available_formats()
        logger.info(f"CV Parser initialized. Supported formats: {list(self.supported_formats.keys())}")
    
    def _get_available_formats(self) -> Dict[str, str]:
        """Get list of actually supported formats based on installed libraries."""
        formats = {'txt': 'Plain Text'}
        
        if PDF_AVAILABLE:
            formats['pdf'] = 'PDF Document'
            logger.info("✅ PDF support enabled")
        else:
            logger.warning("❌ PDF support disabled - install pypdf")
        
        if DOCX_AVAILABLE:
            formats['docx'] = 'Word Document'
            formats['doc'] = 'Word Document (Legacy)'
            logger.info("✅ DOCX support enabled")
        else:
            logger.warning("❌ DOCX support disabled - install python-docx")
        
        if OCR_AVAILABLE:
            formats.update({
                'png': 'Image (PNG)',
                'jpg': 'Image (JPEG)',
                'jpeg': 'Image (JPEG)'
            })
            logger.info("✅ Image OCR support enabled")
        else:
            logger.warning("❌ Image OCR disabled - install pytesseract")
        
        return formats
    
    def parse_cv(self, file_content: bytes, filename: str) -> Optional[str]:
        """
        Parse CV file and extract text.
        
        Args:
            file_content: File content as bytes
            filename: Original filename
        
        Returns:
            Extracted text or None if parsing failed
        """
        # Check file size
        if len(file_content) > self.MAX_FILE_SIZE:
            logger.error(f"File too large: {len(file_content)} bytes (max: {self.MAX_FILE_SIZE})")
            raise ValueError(f"File size exceeds maximum of {self.MAX_FILE_SIZE / 1024 / 1024:.0f} MB")
        
        # Get file extension
        ext = Path(filename).suffix.lower().lstrip('.')
        
        if ext not in self.supported_formats:
            logger.error(f"Unsupported file format: {ext}")
            raise ValueError(f"Unsupported file format: .{ext}. Supported formats: {', '.join(self.supported_formats.keys())}")
        
        # Parse based on file type
        try:
            logger.info(f"Parsing {ext.upper()} file: {filename}")
            
            if ext == 'txt':
                text = self._parse_txt(file_content)
            elif ext == 'pdf':
                text = self._parse_pdf(file_content)
            elif ext in ['docx', 'doc']:
                text = self._parse_docx(file_content)
            elif ext in ['png', 'jpg', 'jpeg']:
                text = self._parse_image(file_content)
            else:
                raise ValueError(f"No parser available for: {ext}")
            
            if text:
                logger.info(f"✅ Successfully extracted {len(text)} characters")
                return text
            else:
                logger.warning("⚠️ No text extracted from file")
                return None
        
        except Exception as e:
            logger.error(f"❌ Error parsing CV: {e}", exc_info=True)
            raise
    
    def _parse_txt(self, content: bytes) -> str:
        """Parse plain text file."""
        try:
            # Try UTF-8 first
            return content.decode('utf-8')
        except UnicodeDecodeError:
            # Fallback to latin-1
            try:
                return content.decode('latin-1')
            except:
                return content.decode('utf-8', errors='ignore')
    
    def _parse_pdf(self, content: bytes) -> Optional[str]:
        """Parse PDF file."""
        if not PDF_AVAILABLE:
            raise ImportError("PDF parsing not available. Please install: pip install pypdf")
        
        try:
            pdf_file = io.BytesIO(content)
            pdf_reader = PdfReader(pdf_file)
            
            logger.info(f"PDF has {len(pdf_reader.pages)} pages")
            
            text_parts = []
            for i, page in enumerate(pdf_reader.pages):
                try:
                    text = page.extract_text()
                    if text and text.strip():
                        text_parts.append(text)
                        logger.debug(f"Extracted {len(text)} chars from page {i+1}")
                except Exception as e:
                    logger.warning(f"Could not extract text from page {i+1}: {e}")
            
            if not text_parts:
                logger.warning("No text extracted from PDF. It might be an image-based PDF.")
                return None
            
            full_text = '\n\n'.join(text_parts)
            return full_text
        
        except Exception as e:
            logger.error(f"Error parsing PDF: {e}", exc_info=True)
            raise ValueError(f"Failed to parse PDF: {str(e)}")
    
    def _parse_docx(self, content: bytes) -> Optional[str]:
        """Parse DOCX file."""
        if not DOCX_AVAILABLE:
            raise ImportError("DOCX parsing not available. Please install: pip install python-docx")
        
        try:
            docx_file = io.BytesIO(content)
            doc = Document(docx_file)
            
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            
            if not text_parts:
                logger.warning("No text extracted from DOCX")
                return None
            
            return '\n'.join(text_parts)
        
        except Exception as e:
            logger.error(f"Error parsing DOCX: {e}", exc_info=True)
            raise ValueError(f"Failed to parse DOCX: {str(e)}")
    
    def _parse_image(self, content: bytes) -> Optional[str]:
        """Parse image file using OCR."""
        if not OCR_AVAILABLE:
            raise ImportError("OCR not available. Please install: pip install pytesseract pillow")
        
        try:
            image = Image.open(io.BytesIO(content))
            text = pytesseract.image_to_string(image)
            return text.strip() if text else None
        
        except Exception as e:
            logger.error(f"Error parsing image: {e}", exc_info=True)
            raise ValueError(f"Failed to parse image: {str(e)}")
    
    def extract_cv_summary(self, cv_text: str, max_length: int = 1000) -> str:
        """
        Extract a summary from CV text for better matching.
        
        Args:
            cv_text: Full CV text
            max_length: Maximum length of summary
        
        Returns:
            CV summary
        """
        if not cv_text:
            return ""
        
        # Take first max_length characters
        summary = cv_text[:max_length]
        
        # Try to end at a sentence boundary
        if len(cv_text) > max_length:
            last_period = summary.rfind('.')
            last_newline = summary.rfind('\n')
            
            cut_point = max(last_period, last_newline)
            if cut_point > max_length * 0.7:  # At least 70% of max_length
                summary = summary[:cut_point + 1]
        
        return summary.strip()


# Singleton instance
_cv_parser = None

def get_cv_parser() -> CVParser:
    """Get CV parser singleton instance."""
    global _cv_parser
    if _cv_parser is None:
        _cv_parser = CVParser()
    return _cv_parser


if __name__ == "__main__":
    # Test the parser
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(levelname)s - %(message)s'
    )
    
    parser = CVParser()
    
    print("=" * 60)
    print("CV Parser Test")
    print("=" * 60)
    print(f"\nSupported formats: {list(parser.supported_formats.keys())}")
    print(f"Max file size: {parser.MAX_FILE_SIZE / 1024 / 1024:.1f} MB")
    
    # Test with a sample text
    sample_cv = b"""
    John Doe
    Software Engineer
    
    EXPERIENCE:
    - 5 years Python development
    - 3 years React/Node.js
    - Cloud deployment (AWS, Azure)
    
    SKILLS:
    Python, JavaScript, Docker, Kubernetes
    """
    
    print("\n" + "=" * 60)
    print("Testing TXT parsing...")
    print("=" * 60)
    
    result = parser.parse_cv(sample_cv, "test.txt")
    print(f"\nParsed text:\n{result}")
    
    summary = parser.extract_cv_summary(result, max_length=200)
    print(f"\nSummary:\n{summary}")
    
    print("\n✅ Test complete!")